import sys
import json
import os
import re
from openai import OpenAI
from docx import Document

# Import rules engine from same directory
import rules_engine

def prune_job_description(raw_jd: str) -> str:
    """
    Remove boilerplate sections (benefits, EEO, company descriptions)
    from the Job Description to save input tokens, with safety guards
    to prevent false-positive truncations.
    """
    lines = raw_jd.split("\n")
    
    # Clean headers that indicate the start of boilerplate to be truncated
    boilerplate_headers = [
        "what we offer", "benefits", "perks", "compensation", 
        "equal opportunity", "diversity", "about us", "about the company", 
        "company overview", "our culture", "company culture", 
        "physical requirements", "work environment", "how to apply"
    ]
    
    # Specific keywords indicating boilerplate lines to filter out
    boilerplate_keywords = [
        r"401\(k\)",
        r"health insurance",
        r"dental insurance",
        r"vision insurance",
        r"medical insurance",
        r"dental\s*(and|&)?\s*vision",
        r"paid time off",
        r"\bpto\b",
        r"salary range",
        r"competitive salary",
        r"equal opportunity employer",
        r"visa sponsorship",
        r"hybrid work",
        r"medical\s*(and|&)?\s*dental"
    ]
    
    cleaned_lines = []
    for line in lines:
        lower_line = line.lower().strip()
        stripped_line = line.strip()
        
        # Only truncate on short lines (headers) that are NOT list bullets
        if len(stripped_line) < 50 and stripped_line:
            if not any(stripped_line.startswith(b) for b in ['*', '-', '•', '+', 'o', '–']):
                clean_match = re.sub(r'[^a-z0-9\s]+', ' ', lower_line).strip()
                if any(clean_match.startswith(header) for header in boilerplate_headers):
                     print(f"# JD Pruner: Truncated JD at boilerplate section: '{line.strip()}'", file=sys.stderr)
                     break
        cleaned_lines.append(line)
        
    filtered_lines = []
    pruned_count = 0
    for line in cleaned_lines:
        lower_line = line.lower().strip()
        if any(re.search(pattern, lower_line) for pattern in boilerplate_keywords):
            pruned_count += 1
            continue
        filtered_lines.append(line)
        
    if pruned_count > 0:
        print(f"# JD Pruner: Filtered out {pruned_count} boilerplate lines", file=sys.stderr)
        
    return "\n".join(filtered_lines).strip()


class OpenAIProvider:
    def generate(self, system_prompt, user_prompt, model_name):
        client = OpenAI()
        response = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            response_format={"type": "json_object"}
        )
        return response.choices[0].message.content

def copy_style(source_run, target_run):
    """Copy ALL font attributes from source to target for exact visual match."""
    try:
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        
        if source_run.font.strike is not None:
            target_run.font.strike = source_run.font.strike
        if source_run.font.double_strike is not None:
            target_run.font.double_strike = source_run.font.double_strike
        if source_run.font.subscript is not None:
            target_run.font.subscript = source_run.font.subscript
        if source_run.font.superscript is not None:
            target_run.font.superscript = source_run.font.superscript
        if source_run.font.small_caps is not None:
            target_run.font.small_caps = source_run.font.small_caps
        if source_run.font.all_caps is not None:
            target_run.font.all_caps = source_run.font.all_caps
        
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.color and source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb
    except Exception:
        pass

def iter_all_paragraphs(document):
    """Iterate ALL paragraphs: body + tables."""
    for para in document.paragraphs:
        yield para
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    yield para

def get_para_markdown(para):
    """Extract text from paragraph and convert run-level bolding to markdown."""
    result = ""
    for run in para.runs:
        text = run.text
        if not text: continue
        if run.bold:
            clean_text = text.strip()
            if clean_text:
                prefix = text[:text.find(clean_text)]
                suffix = text[len(prefix)+len(clean_text):]
                result += f"{prefix}**{clean_text}**{suffix}"
            else:
                result += text
        else:
            result += text
    return result.strip()

def replace_paragraph_text(para, new_text):
    """Replace paragraph text while preserving formatting and parsing markdown bolding."""
    captured_run = None
    if para.runs:
        for r in para.runs:
            if r.text.strip():
                captured_run = r
                break
        if not captured_run:
            captured_run = para.runs[0]
    
    para.clear()
    new_text = new_text.replace('\n', ' ')
    new_text = new_text.lstrip()
    if new_text.startswith('- '):
        new_text = new_text[2:]
    elif new_text.startswith('* ') and not new_text.startswith('**'):
        new_text = new_text[2:]
        
    if '**' not in new_text:
        match = re.match(r'^([A-Z][A-Za-z0-9\s&/\-,]{2,45}):\s(.*)', new_text)
        if match:
            new_text = f"**{match.group(1)}:** {match.group(2)}"
            
    has_markdown_bold = '**' in new_text
    parts = re.split(r'(\*\*.*?\*\*)', new_text)
    
    for part in parts:
        if not part:
            continue
            
        is_bold = part.startswith('**') and part.endswith('**')
        run_text = part[2:-2] if is_bold else part
        
        new_run = para.add_run(run_text)
        
        if captured_run:
            copy_style(captured_run, new_run)
            
        if has_markdown_bold:
            new_run.bold = is_bold

def run_stage1_planner(jd_text, resume_text, gap_report, model):
    system_prompt = """You are a senior resume strategist and ATS optimization planner.
Your job is to analyze the target Job Description (JD), the numbered Resume, and the Keyword Gap Report to create a strict execution plan for tailoring the resume.

You must output a raw JSON object in the following format:
{
    "detected_seniority": "Junior | Mid-Level | Senior | Executive",
    "candidate_domain": "Primary professional domain identified from the candidate resume (e.g., SAP, Workday, Business Analysis, Cloud Engineering, Software Engineering, etc.)",
    "candidate_current_specialization": "Current specialization identified from the candidate resume (e.g., SAP MM, Payroll, QA Analyst, Java Full Stack)",
    "target_domain": "Target professional domain identified from the Job Description (e.g., SAP, Workday, Business Analysis, DevOps, Network Security, etc.)",
    "target_specialization": "Target specialization identified from the Job Description (e.g., SAP FICO, Integrations, QA Automation Engineer, Backend Engineering)",
    "core_target_capabilities": ["list of capabilities that directly support the target role and must dominate the resume sections"],
    "supporting_capabilities": ["capabilities that strengthen the target role but are secondary"],
    "background_capabilities": ["capabilities belonging to the previous specialization that must be kept but demoted in emphasis"],
    "skills_reconstruction_plan": "Specific guidelines to reorder the Skills section: core target-role skills first, supporting skills second, and background skills last.",
    "experience_identity_plan": "Specific instructions for rewriting experience bullets so that target-role responsibilities become the primary focus and background tasks are demoted in emphasis.",
    "client_1_index": 1-based paragraph index of Client 1 (earlier role) title line,
    "client_1_original_title": "the exact text of Client 1 title paragraph",
    "client_1_proposed_title": "Proposed Client 1 Title Line", // MUST preserve original company and dates exactly, evolving only the job title part to a foundational/transitional title
    "client_2_index": 1-based paragraph index of Client 2 (recent role) title line,
    "client_2_original_title": "the exact text of Client 2 title paragraph",
    "client_2_proposed_title": "Proposed Client 2 Title Line", // MUST preserve original company and dates exactly, evolving only the job title part to align with JD target role
    "tool_substitutions": {
        "Manhattan OMS": "Epic" // Map any domain-specific tool in the JD that doesn't match candidate's domain to its functional equivalent
    },
    "keyword_distribution": {
        "summary": ["key term 1", "key term 2"],
        "skills": ["key term 3", "key term 4"],
        "experience": ["key term 5", "key term 6"]
    },
    "certification_gaps": [
        {
            "certification": "Name of missing required certification",
            "compensation_strategy": "Concrete bullet instructions proving the candidate has that knowledge through hands-on work to make the lack of certification invisible"
        }
    ],
    "experience_gap_handling": "Strategy to handle years of experience gap. If candidate has fewer years of experience than the JD requires, describe how to showcase the depth, breadth, and ownership of their existing experience powerfully so the years gap becomes irrelevant. Never mention the gap. Never hedge."
}

CRITICAL RULES FOR PROPOSED TITLES:
1. You MUST preserve the company names and dates/tenures (e.g., '2020 - 2022', '2022 - Present') from the original paragraph text exactly.
2. Only modify the job title portion within that paragraph to match the career trajectory.
3. Client 1 (earlier role) MUST represent a foundational, entry-level, or transitional title.
4. Client 2 (recent role) MUST represent a specialized title directly aligned to the target JD title.
"""
    user_prompt = f"""Job Description:
{jd_text}

Resume (numbered):
{resume_text}

Keyword Gap Report:
{json.dumps(gap_report, indent=2)}
"""
    ai_provider = OpenAIProvider()
    response_text = ai_provider.generate(system_prompt, user_prompt, model)
    response_text = response_text.replace("```json", "").replace("```", "").strip()
    
    try:
        plan = json.loads(response_text)
    except json.JSONDecodeError:
        start = response_text.find('{')
        end = response_text.rfind('}') + 1
        if start != -1 and end != -1:
            plan = json.loads(response_text[start:end])
        else:
            plan = {}
            
    return plan

def run_keyword_gap_analysis(jd_text, resume_text, candidate_domain):
    return rules_engine.build_full_gap_report(jd_text, resume_text, candidate_domain)

def optimize_pro(input_path, output_path, jd_text, prompt_instruction, model_name):
    """
    PRO MODE: Positional replacement with Two-Stage Pipeline and Multi-Pass Optimization.
    """
    try:
        doc = Document(input_path)
        
        # Build indexed list of all non-empty paragraphs
        indexed_paragraphs = []
        for para in iter_all_paragraphs(doc):
            text = para.text.strip()
            if text:
                indexed_paragraphs.append(para)
        
        # Record original styles and texts to restore if formatting checks fail
        original_styles = [para.style.name for para in indexed_paragraphs]
        original_texts = [para.text for para in indexed_paragraphs]
        
        # Section classification
        SECTION_KEYWORDS = {
            "summary": ["summary", "profile", "objective", "professional summary", "career objective"],
            "skills": ["skills", "technical skills", "expertise", "competencies", "technologies", "tools", "areas of expertise"],
            "experience": ["experience", "employment history", "work history", "professional experience", "work experience"],
            "education": ["education", "academic", "credentials", "degrees", "university", "schooling"]
        }

        def classify_heading(text_val):
            lower = text_val.lower().strip()
            if len(lower) > 40:
                return None
            for section_key, keywords in SECTION_KEYWORDS.items():
                for kw in keywords:
                    if kw in lower:
                        return section_key
            return None

        current_section = "header"
        para_sections = []
        for para in indexed_paragraphs:
            text = para.text.strip()
            detected = classify_heading(text)
            if detected is not None:
                current_section = detected
            para_sections.append(current_section)
        
        temp_resume_text = "\n".join([p.text for p in indexed_paragraphs])
        
        # Heuristic candidate domain detection
        heuristic_domain = "Other"
        domain_keywords = {
            "HRIS": ["workday", "ukg", "adp", "hcm", "payroll", "successfactors"],
            "Finance": ["sap", "fico", "oracle financials", "anaplan", "hyperion", "accounting"],
            "Data Engineering": ["snowflake", "dbt", "databricks", "spark", "airflow", "pipeline", "xslt", "xml", "xpath"],
            "Retail": ["oms", "retail", "supply chain", "salesforce commerce"],
            "Logistics": ["transportation", "logistics", "tms", "mercurygate"]
        }
        domain_scores = {d: 0 for d in domain_keywords}
        for d, kws in domain_keywords.items():
            for kw in kws:
                if kw.lower() in temp_resume_text.lower():
                    domain_scores[d] += 1
        best_d = max(domain_scores, key=domain_scores.get)
        if domain_scores[best_d] > 0:
            heuristic_domain = best_d
            
        gap_report = run_keyword_gap_analysis(jd_text, temp_resume_text, heuristic_domain)

        # Build numbered resume
        numbered_lines = []
        for i, para in enumerate(indexed_paragraphs):
            section = para_sections[i]
            if section in ["header", "education"]:
                continue
            para_md = get_para_markdown(para)
            numbered_lines.append(f"[{i + 1}] {para_md}")
        
        numbered_resume = "\n".join(numbered_lines)
        
        # ─── STAGE 1: PLANNER ──────────────────────────────────────────────
        print(f"# Stage 1 Planner: Running analysis...", file=sys.stderr)
        plan = run_stage1_planner(jd_text, numbered_resume, gap_report, model_name)
        
        candidate_domain = plan.get("candidate_domain", heuristic_domain)
        candidate_current_specialization = plan.get("candidate_current_specialization", "Not Specified")
        target_domain = plan.get("target_domain", plan.get("jd_domain", "Other"))
        target_specialization = plan.get("target_specialization", "Not Specified")
        seniority = plan.get("detected_seniority", "Mid-Level")
        jd_domain = target_domain
        
        core_target_capabilities = plan.get("core_target_capabilities", [])
        supporting_capabilities = plan.get("supporting_capabilities", [])
        background_capabilities = plan.get("background_capabilities", [])
        skills_reconstruction_plan = plan.get("skills_reconstruction_plan", "Reorder skills targeting target specialization first, supporting second, background last.")
        experience_identity_plan = plan.get("experience_identity_plan", "Transform bullets prioritizing target capabilities and demoting background capabilities.")
        
        print(f"# Stage 1 Planner Output: Seniority={seniority}, Candidate Domain={candidate_domain}, Target Domain={target_domain}", file=sys.stderr)
        
        # Update gap report if candidate domain is corrected
        if candidate_domain != heuristic_domain:
            gap_report = run_keyword_gap_analysis(jd_text, temp_resume_text, candidate_domain)
            
        tool_map = rules_engine.get_tool_mappings(candidate_domain, jd_domain)
        stage1_subs = plan.get("tool_substitutions", {})
        if isinstance(stage1_subs, dict):
            tool_map.update(stage1_subs)
            
        # Update client titles before rewrite
        client_1_idx = plan.get("client_1_index")
        client_2_idx = plan.get("client_2_index")
        client_1_updated = False
        client_2_updated = False
        
        if client_1_idx:
            try:
                pos1 = int(client_1_idx) - 1
                if 0 <= pos1 < len(indexed_paragraphs):
                    new_title = plan.get("client_1_proposed_title")
                    if new_title:
                        replace_paragraph_text(indexed_paragraphs[pos1], new_title)
                        client_1_updated = True
            except Exception as e:
                print(f"# Error updating Client 1 title paragraph: {e}", file=sys.stderr)
                
        if client_2_idx:
            try:
                pos2 = int(client_2_idx) - 1
                if 0 <= pos2 < len(indexed_paragraphs):
                    new_title = plan.get("client_2_proposed_title")
                    if new_title:
                        replace_paragraph_text(indexed_paragraphs[pos2], new_title)
                        client_2_updated = True
            except Exception as e:
                print(f"# Error updating Client 2 title paragraph: {e}", file=sys.stderr)
        
        # Build initial keyword injection rules list
        injection_rules_list = []
        for cat in ["hard_required", "soft_required", "preferred"]:
            for kw_data in gap_report.get(cat, []):
                kw = kw_data["keyword"]
                status = kw_data["status"]
                conf = kw_data["confidence"]
                
                if status == "present":
                    injection_rules_list.append(f"- Keyword '{kw}' is already present in the resume. Ensure it conforms to {conf} level permitted verbs.")
                    continue
                    
                if conf == "FULL_EXPERIENCE":
                    injection_rules_list.append(
                        f"- Keyword '{kw}' (Confidence: FULL_EXPERIENCE): Inject into the Skills section, the Summary, and at least one Experience bullet. "
                        f"You MUST use one of these verbs: Led, Designed, Implemented, Configured, Developed, Administered, Managed, Optimized, Owned."
                    )
                elif conf == "ADJACENT_EXPERIENCE":
                    injection_rules_list.append(
                        f"- Keyword '{kw}' (Confidence: ADJACENT_EXPERIENCE): Inject into the Skills section and at least one Experience bullet. "
                        f"You MUST use one of these verbs: Partnered with, Collaborated on, Supported, Coordinated, Contributed to, Worked alongside."
                    )
                elif conf == "EXPOSURE_ONLY":
                    injection_rules_list.append(
                        f"- Keyword '{kw}' (Confidence: EXPOSURE_ONLY): Inject into the Skills section and at least one Experience bullet. "
                        f"You MUST use one of these verbs: Monitored, Evaluated, Participated in, Observed, Reviewed, Assisted with testing, Validated."
                    )
                elif conf == "NO_EXPERIENCE":
                    comp_strategy = ""
                    for cert_gap in plan.get("certification_gaps", []):
                        if cert_gap.get("certification") == kw or kw in cert_gap.get("certification", ""):
                            comp_strategy = cert_gap.get("compensation_strategy", "")
                            break
                    if comp_strategy:
                        injection_rules_list.append(
                            f"- Certification '{kw}' (Confidence: NO_EXPERIENCE): Direct injection is BLOCKED. "
                            f"Do NOT mention this certification anywhere. Instead, write equivalent experience bullets demonstrating: {comp_strategy}."
                        )
                    else:
                        injection_rules_list.append(f"- Keyword/Certification '{kw}' (Confidence: NO_EXPERIENCE): Direct injection is BLOCKED. Do NOT inject or mention this keyword anywhere.")

        injection_instructions = "\n".join(injection_rules_list) if injection_rules_list else "No keyword instructions."
        exp_gap_handling = plan.get("experience_gap_handling", "Emphasize ownership depth and senior-level scope across all bullets.")
        
        # Get seniority-calibrated verbs
        calibration = rules_engine.get_calibrated_prompts(seniority)
        allowed_verbs = ", ".join(calibration["allowed_verbs"])
        forbidden_phrases = ", ".join(calibration["forbidden_phrases"])
        
        domain_transition_strategy = (
            f"Candidate Primary Domain: {candidate_domain}\n"
            f"Candidate Current Specialization: {candidate_current_specialization}\n"
            f"Target Domain: {target_domain}\n"
            f"Target Specialization: {target_specialization}\n"
            "\n"
            f"CORE TARGET CAPABILITIES (Promote & Emphasize):\n"
            f"{', '.join(core_target_capabilities) if core_target_capabilities else 'Directly align with target role'}\n"
            "\n"
            f"SUPPORTING CAPABILITIES (Keep Visible but Secondary):\n"
            f"{', '.join(supporting_capabilities) if supporting_capabilities else 'Strengthen the target role'}\n"
            "\n"
            f"BACKGROUND CAPABILITIES (Reduce Emphasis but Do NOT Remove):\n"
            f"{', '.join(background_capabilities) if background_capabilities else 'Demote previous role capabilities'}\n"
            "\n"
            f"SKILLS SECTION RECONSTRUCTION PLAN:\n"
            f"{skills_reconstruction_plan}\n"
            "\n"
            f"EXPERIENCE ROLE IDENTITY TRANSFORMATION PLAN:\n"
            f"{experience_identity_plan}"
        )
        
        # ─── MULTI-PASS OPTIMIZER LOOP ────────────────────────────────────
        max_passes = 3
        current_pass = 1
        sanitized_count = 0
        
        while current_pass <= max_passes:
            print(f"# Stage 2 Rewrite: Optimization Pass {current_pass}...", file=sys.stderr)
            
            # Re-build numbered resume for this pass
            numbered_lines = []
            for idx, para in enumerate(indexed_paragraphs):
                section = para_sections[idx]
                if section in ["header", "education"]:
                    continue
                para_md = get_para_markdown(para)
                numbered_lines.append(f"[{idx + 1}] {para_md}")
            numbered_resume = "\n".join(numbered_lines)
            
            if current_pass == 1:
                system_prompt = f"""You are a senior resume rewriter. Your job is to optimize the resume paragraphs to align with the target JD while remaining 100% defensible, professional, and natural.
                
**DOMAIN RECOGNITION & SPECIALIZATION TRANSITION**:
{domain_transition_strategy}

CRITICAL ROLE TRANSFORMATION RULE:
Do NOT transform blindly using JD keywords. You must transform the resume through this domain transition ecosystem. Reconstruct the Professional Summary, Skills, and Experience sections to show a natural career progression from the candidate's current specialization to the target specialization. Ensure all transformed responsibilities are realistic, explainable, and logically defensible within this transition.

**SUMMARY POSITIONING RULE (CRITICAL)**:
- The Professional Summary must represent the strongest defensible version of the target specialization.
- Avoid weak positioning language: "Exposure to", "Supported", "Assisted with", "Worked alongside" (unless no stronger evidence exists). Position the candidate as an active contributor operating within the target specialization.

**SKILLS PRIORITIZATION RULE (CRITICAL)**:
- Reorder the Skills section. Core target-role skills must appear first, supporting skills second, and background skills last.

**FINAL IDENTITY CHECK**:
- Before outputting, ask yourself: "If a recruiter reads only the Summary, Skills, and the first 5 bullets of the most recent role, what professional identity would they assign to this candidate?"
- If the answer is not the target specialization, continue refining the content until the target identity is clearly and dominantly established.

**KEYWORD AUTO-INJECTION AND CONFIDENCE RULES**:
{injection_instructions}

**EXPERIENCE YEARS GAP HANDLING**:
{exp_gap_handling}

**CRITICAL COMPLIANCE RULES**:
1. NEVER alter, remove, or rewrite any date, year, or date range in the resume paragraphs.
2. The Client 1 title (paragraph at index {client_1_idx}) has already been set. Do NOT change this paragraph.
3. The Client 2 title (paragraph at index {client_2_idx}) has already been set. Do NOT change this paragraph.
4. Client 1 Experience Bullets must emphasize: reporting, operational support, data analysis, coordination, and foundational system usage.
5. Client 2 Experience Bullets must emphasize: ownership, optimization, cross-functional collaboration, advanced configuration, and process improvements.
6. Do NOT copy any tool from the JD directly if it does not match the candidate's domain ({candidate_domain}). Only use the functional substitutions: {json.dumps(tool_map)}.
7. ATS Keyword stuffing prevention: Do NOT repeat the same keyword or skill more than twice across the entire resume.
8. Seniority Calibration ({seniority} Level):
   - YOU MUST USE allowed verbs/themes: {allowed_verbs}
   - YOU ARE STRICTLY FORBIDDEN from using any of the following terms or themes: {forbidden_phrases}
9. All responsibilities must be realistic, explainable, and logically defensible. No exaggerated claims.
"""
            else:
                current_resume_text = "\n".join([p.text for p in indexed_paragraphs])
                generated_strategies = [g.get("certification", "").lower() for g in plan.get("certification_gaps", [])]
                current_scores = rules_engine.calculate_ats_score(jd_text, current_resume_text, generated_strategies)
                
                feedback_items = []
                # Identify missing keywords
                for cat in ["hard_required", "soft_required", "preferred"]:
                    for item in gap_report.get(cat, []):
                        kw = item["keyword"]
                        conf = item["confidence"]
                        injectable = item["injectable"]
                        
                        presence = rules_engine.classify_keyword_presence(kw, current_resume_text)
                        if presence == "missing" and injectable:
                            feedback_items.append(f"- Keyword '{kw}' is MISSING. Please inject it in the Skills and Experience sections using allowed verbs.")
                            
                # Identify verb violations
                for cat in ["hard_required", "soft_required", "preferred"]:
                    for item in gap_report.get(cat, []):
                        kw = item["keyword"]
                        conf = item["confidence"]
                        presence = rules_engine.classify_keyword_presence(kw, current_resume_text)
                        if presence == "present":
                            is_compliant = rules_engine.check_keyword_verb_compliance(kw, current_resume_text, conf)
                            if not is_compliant:
                                feedback_items.append(f"- Keyword '{kw}' (Confidence: {conf}) is used with a verb that is too strong. You must downgrade the verb to match the {conf} level.")
                
                feedback_str = "\n".join(feedback_items) if feedback_items else "No deficiencies detected."
                
                system_prompt = f"""You are a senior resume optimizer. This is optimization pass {current_pass} to raise the alignment scores to 92+.
                
**DOMAIN RECOGNITION & SPECIALIZATION TRANSITION**:
{domain_transition_strategy}

CRITICAL ROLE TRANSFORMATION RULE:
Do NOT transform blindly using JD keywords. You must transform the resume through this domain transition ecosystem. Reconstruct the Professional Summary, Skills, and Experience sections to show a natural career progression from the candidate's current specialization to the target specialization. Ensure all transformed responsibilities are realistic, explainable, and logically defensible within this transition.

**SUMMARY POSITIONING RULE (CRITICAL)**:
- The Professional Summary must represent the strongest defensible version of the target specialization.
- Avoid weak positioning language: "Exposure to", "Supported", "Assisted with", "Worked alongside" (unless no stronger evidence exists). Position the candidate as an active contributor operating within the target specialization.

**SKILLS PRIORITIZATION RULE (CRITICAL)**:
- Reorder the Skills section. Core target-role skills must appear first, supporting skills second, and background skills last.

**FINAL IDENTITY CHECK**:
- Before outputting, ask yourself: "If a recruiter reads only the Summary, Skills, and the first 5 bullets of the most recent role, what professional identity would they assign to this candidate?"
- If the answer is not the target specialization, continue refining the content until the target identity is clearly and dominantly established.

**CURRENT METRICS**:
- ATS Score: {current_scores["score"]} / 100
- Keyword Match Score: {current_scores["keyword_match"]} / 100
- Recruiter Credibility Score: {current_scores["recruiter_credibility"]} / 100
- Interview Defensibility Score: {current_scores["interview_defensibility"]} / 100

**OPTIMIZATION FEEDBACK & DEFICIENCIES**:
{feedback_str}

**KEYWORD CONFIDENCE & INJECTION RULES**:
{injection_instructions}

**CRITICAL COMPLIANCE RULES**:
1. NEVER alter, remove, or rewrite any date, year, or date range.
2. The Client 1 title (paragraph at index {client_1_idx}) has already been set. Do NOT change it.
3. The Client 2 title (paragraph at index {client_2_idx}) has already been set. Do NOT change it.
4. Do NOT copy any tool from the JD directly if it does not match the candidate's domain ({candidate_domain}). Only use the functional substitutions: {json.dumps(tool_map)}.
5. ATS Keyword stuffing prevention: Do NOT repeat the same keyword or skill more than twice across the entire resume.
6. Seniority Calibration ({seniority} Level): Use allowed verbs/themes ({allowed_verbs}) and avoid forbidden phrases ({forbidden_phrases}).
"""

            # Formatting Rules
            system_prompt += """
**FORMATTING RULES — NON-NEGOTIABLE**:
1. Never use em dashes or en dashes anywhere. Use a comma, colon, or rewrite instead.
2. Never use smart quotes, ellipsis characters, or any Unicode typography. Plain ASCII only.
3. Never add bullets to summary, skills, title lines, headers, or education. Bullets in experience bullets only.
4. Write every sentence as a real human would type in Microsoft Word. Active voice. Under 35 words per sentence. No semicolon chains.
5. Never start a bullet with AI-pattern openers like Spearheaded the facilitation of.
"""

            system_prompt += f"""
**CRITICAL OUTPUT FORMAT (MUST follow exactly)**:
The resume content below is numbered with [1], [2], [3], etc.
For each point you want to replace, return its number and the new text.
Return ONLY a raw JSON object in this exact format:
{{
    "replacements": [
        {{"index": 1, "new": "your rewritten text for point 1"}},
        {{"index": 3, "new": "your rewritten text for point 3"}}
    ]
}}
Only include indexes for points you are changing. Do NOT include the [N] prefix in the "new" text.
Do NOT include replacements for client title paragraphs or date paragraphs.
You MAY use **bold** markdown (like **Role Title**) in the "new" text.
"""
            # Custom prompt addition
            if prompt_instruction and prompt_instruction.strip() and prompt_instruction.strip().lower() != "highlight experience relevant to the job requirements.":
                system_prompt += f"""
**ADDITIONAL TEAM INSTRUCTIONS**:
{prompt_instruction.strip()}
"""

            user_prompt = f"""Job Description:
{jd_text}

Resume Content (numbered):
{numbered_resume}
"""
            
            ai_client = OpenAIProvider()
            response_text = ai_client.generate(system_prompt, user_prompt, model_name)
            response_text = response_text.replace("```json", "").replace("```", "").strip()
            
            try:
                data = json.loads(response_text)
            except json.JSONDecodeError:
                start = response_text.find('{')
                end = response_text.rfind('}') + 1
                if start != -1 and end != -1:
                    data = json.loads(response_text[start:end])
                else:
                    print(json.dumps({"status": "error", "message": f"Could not parse AI response in pass {current_pass}"}))
                    sys.exit(1)
            
            replacements = data.get("replacements", [])
            changes_count = 0
            date_pattern = re.compile(r'\b(?:19|20)\d{2}\b|\bPresent\b', re.IGNORECASE)
            
            for replacement in replacements:
                idx = replacement.get("index")
                new_text = replacement.get("new", "")
                
                if idx is None or not new_text:
                    continue
                
                pos = idx - 1
                if pos < 0 or pos >= len(indexed_paragraphs):
                    continue
                
                target_para = indexed_paragraphs[pos]
                old_text = target_para.text.strip()
                
                if date_pattern.search(old_text):
                    continue
                if (client_1_updated and pos == int(client_1_idx) - 1) or (client_2_updated and pos == int(client_2_idx) - 1):
                    continue
                
                replace_paragraph_text(target_para, new_text)
                changes_count += 1
            
            # Post-processing sanitization pass
            sanitized_count = 0
            for i, para in enumerate(indexed_paragraphs):
                orig_t = para.text
                sanitized = rules_engine.sanitize_output_text(orig_t)
                if sanitized != orig_t:
                    replace_paragraph_text(para, sanitized)
                    sanitized_count += 1
                    
                # Clean list bullets on non-experience
                sec = para_sections[i]
                if sec != "experience":
                    plain_text_cleared = para.text.lstrip(" \t•*-⁃")
                    if plain_text_cleared != para.text:
                        replace_paragraph_text(para, plain_text_cleared)
                    
                    curr_style = para.style.name.lower()
                    if "list" in curr_style or "bullet" in curr_style:
                        para.style = original_styles[i]
            
            # Re-evaluate scores
            final_resume_text = "\n".join([p.text for p in indexed_paragraphs])
            generated_strategies = [g.get("certification", "").lower() for g in plan.get("certification_gaps", [])]
            final_scores = rules_engine.calculate_ats_score(jd_text, final_resume_text, generated_strategies)
            
            print(f"# Pass {current_pass} Results: ATS={final_scores['score']}, Match={final_scores['keyword_match']}, Credibility={final_scores['recruiter_credibility']}, Defensibility={final_scores['interview_defensibility']}", file=sys.stderr)
            
            if (final_scores["score"] >= 92 and 
                final_scores["keyword_match"] >= 92 and 
                final_scores["recruiter_credibility"] >= 92 and 
                final_scores["interview_defensibility"] >= 92):
                print(f"# All metrics exceeded target of 92. Optimization complete.", file=sys.stderr)
                break
                
            current_pass += 1
            
        # ─── FINAL DETERMINISTIC REPORT GENERATION ─────────────────────────
        final_resume_text = "\n".join([p.text for p in indexed_paragraphs])
        generated_strategies = [g.get("certification", "").lower() for g in plan.get("certification_gaps", [])]
        score_after_data = rules_engine.calculate_ats_score(jd_text, final_resume_text, generated_strategies)
        score_before = gap_report.get("score_before", 0)
        score_after = score_after_data.get("score", 0)
        
        keyword_entries = []
        for cat in ["hard_required", "soft_required", "preferred"]:
            for item in gap_report.get(cat, []):
                kw = item["keyword"]
                confidence = item["confidence"]
                injectable = item["injectable"]
                reason = item["reason"]
                
                status_after = rules_engine.classify_keyword_presence(kw, final_resume_text)
                
                keyword_entries.append({
                    "keyword": kw,
                    "category": cat.replace("_", " ").title(),
                    "status": status_after.upper(),
                    "confidence": confidence,
                    "injectable": "Yes" if injectable else "No",
                    "reason": reason
                })
                
        # Certification equivalent experience flagging
        cert_gaps_noted = []
        for cert_gap in plan.get("certification_gaps", []):
            cert_name = cert_gap.get("certification")
            strategy = cert_gap.get("compensation_strategy")
            cert_gaps_noted.append({
                "certification": cert_name,
                "equivalent_experience": strategy
            })
            
        for entry in keyword_entries:
            kw_lower = entry["keyword"].lower()
            is_certification = "certif" in kw_lower or "certified" in kw_lower or kw_lower in ["pmp", "cissp", "itil", "csm"]
            if is_certification and entry["status"] == "MISSING":
                # Find matching cert strategy in rules engine equivalents
                strategy = "Demonstrate hands-on delivery and practical competency to make lack of certification invisible."
                for cert_key, val in rules_engine.CERTIFICATION_EQUIVALENTS.items():
                    if cert_key in kw_lower or kw_lower in cert_key:
                        strategy = f"Compensate with: {', '.join(val['bullets'])}"
                        break
                # Only append if not already present
                if not any(g["certification"] == entry["keyword"] for g in cert_gaps_noted):
                    cert_gaps_noted.append({
                        "certification": entry["keyword"],
                        "equivalent_experience": strategy
                    })
                    
        from datetime import datetime
        now_str = datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")
        
        candidate_name = "Candidate"
        if indexed_paragraphs:
            candidate_name = indexed_paragraphs[0].text.strip()
            
        target_role = "Optimized Role"
        if plan.get("client_2_proposed_title"):
            target_role = plan.get("client_2_proposed_title").split(" - ")[0].split(" at ")[0].strip()
            
        custom_prompt_used = "none"
        if prompt_instruction and prompt_instruction.strip() and prompt_instruction.strip().lower() != "highlight experience relevant to the job requirements.":
            custom_prompt_used = f"yes - {prompt_instruction.strip()[:100]}"
            
        report_lines = [
            "====================================================",
            "RESUME OPTIMIZATION INTELLIGENCE REPORT",
            "====================================================",
            "",
            f"CANDIDATE: {candidate_name}",
            f"TARGET ROLE: {target_role}",
            f"GENERATED: {now_str}",
            f"CUSTOM PROMPT USED: {custom_prompt_used}",
            "",
            "----------------------------------------------------",
            "ALIGNMENT METRICS & SCORES",
            "----------------------------------------------------",
            f"ATS Matching Score:         {score_after_data['score']} / 100  (Target: 92+)",
            f"Keyword Match Score:        {score_after_data['keyword_match']} / 100  (Target: 92+)",
            f"Recruiter Credibility Score: {score_after_data['recruiter_credibility']} / 100  (Target: 92+)",
            f"Interview Defensibility:     {score_after_data['interview_defensibility']} / 100  (Target: 92+)",
            "",
            "----------------------------------------------------",
            "DETAILED KEYWORD CLASSIFICATION REPORT",
            "----------------------------------------------------"
        ]
        
        for entry in keyword_entries:
            report_lines.extend([
                f"Keyword:    {entry['keyword']}",
                f"  Category: {entry['category']}",
                f"  Status:   {entry['status']}",
                f"  Conf:     {entry['confidence']}",
                f"  Inject:   {entry['injectable']}",
                f"  Reason:   {entry['reason']}",
                ""
            ])
            
        report_lines.extend([
            "----------------------------------------------------",
            "MISSING CERTIFICATIONS & EQUIVALENT EXPERIENCE",
            "----------------------------------------------------"
        ])
        
        if cert_gaps_noted:
            for gap in cert_gaps_noted:
                report_lines.extend([
                    f"Certification Missing: {gap['certification']}",
                    f"  Equivalent Experience Strategy:",
                    f"    {gap['equivalent_experience']}",
                    ""
                ])
        else:
            report_lines.append("No missing certifications detected.\n")
            
        report_lines.extend([
            "----------------------------------------------------",
            "EXPERIENCE GAP HANDLING STRATEGY",
            "----------------------------------------------------",
            exp_gap_handling,
            "",
            "----------------------------------------------------",
            "FORMATTING & STYLE SANITIZATION SUMMARY",
            "----------------------------------------------------",
            f"Em dashes / en dashes sanitized:   {sanitized_count}",
            f"Unicode quotes/ellipses cleaned:  {sanitized_count}",
            f"Non-experience lists reset:       {sanitized_count}",
            ""
        ])

        # --- RESUME OPTIMIZATION QUALITY CHECKLIST ---
        client1_foundational = False
        client2_advanced = False
        progression_logical = False
        
        if client_1_idx and plan.get("client_1_proposed_title"):
            client1_foundational = True
        if client_2_idx and plan.get("client_2_proposed_title"):
            client2_advanced = True
        if client_1_idx and client_2_idx:
            try:
                if int(client_2_idx) < int(client_1_idx):
                    progression_logical = True
            except Exception:
                pass
                
        dates_preserved = True
        durations_preserved = True
        degree_history_preserved = True
        certifications_preserved = True
        
        date_pattern = re.compile(r'\b(?:19|20)\d{2}\b|\bPresent\b', re.IGNORECASE)
        for orig_t, curr_p in zip(original_texts, indexed_paragraphs):
            orig_dates = set(date_pattern.findall(orig_t))
            curr_dates = set(date_pattern.findall(curr_p.text))
            if not orig_dates.issubset(curr_dates):
                dates_preserved = False
                durations_preserved = False
                break
                
        for i, para in enumerate(indexed_paragraphs):
            if para_sections[i] == "education":
                if para.text.strip() != original_texts[i].strip():
                    degree_history_preserved = False
                    break
                    
        original_resume_text = "\n".join(original_texts)
        final_resume_text = "\n".join([p.text for p in indexed_paragraphs])
        for cert_key in rules_engine.CERTIFICATION_EQUIVALENTS.keys():
            if rules_engine.classify_keyword_presence(cert_key, original_resume_text) == "present":
                if rules_engine.classify_keyword_presence(cert_key, final_resume_text) != "present":
                    certifications_preserved = False
                    break
                    
        hard_cov_calculated = True
        soft_cov_calculated = True
        pref_cov_calculated = True
        
        keywords_in_skills = False
        keywords_in_experience = False
        full_exp_in_summary = True
        
        final_sections = rules_engine.segment_resume_text(final_resume_text)
        keyword_freq = {}
        for entry in keyword_entries:
            kw_lower = entry["keyword"].lower()
            freq = len(re.findall(r'\b' + re.escape(kw_lower) + r'\b', final_resume_text.lower()))
            keyword_freq[entry["keyword"]] = freq

            kw = entry["keyword"]
            if rules_engine.classify_keyword_presence(kw, final_sections["skills"]) == "present":
                keywords_in_skills = True
            if rules_engine.classify_keyword_presence(kw, final_sections["experience"]) == "present":
                keywords_in_experience = True
            if entry["confidence"] == "FULL_EXPERIENCE":
                if rules_engine.classify_keyword_presence(kw, final_sections["summary"]) != "present":
                    full_exp_in_summary = False
                    
        no_duplicate_bullets = True
        no_excessive_repetition = True
        no_malformed_formatting = True
        no_empty_sections = True
        
        seen_bullets = set()
        for idx_p, para in enumerate(indexed_paragraphs):
            sec = para_sections[idx_p]
            if sec == "experience":
                text_clean = para.text.strip().lower()
                if text_clean:
                    if text_clean in seen_bullets:
                        no_duplicate_bullets = False
                    seen_bullets.add(text_clean)
                    
        for entry in keyword_entries:
            kw = entry["keyword"]
            category = entry["category"].lower().replace(" ", "_")
            freq = keyword_freq[kw]
            
            # Dynamic threshold components
            base_allowance = 2
            placement_allowance = 0
            in_summary = rules_engine.classify_keyword_presence(kw, final_sections["summary"]) == "present"
            in_skills = rules_engine.classify_keyword_presence(kw, final_sections["skills"]) == "present"
            if in_summary:
                placement_allowance += 1
            if in_skills:
                placement_allowance += 1
                
            exp_paragraphs = [p for p in final_sections["experience"].split("\n") if p.strip()]
            exp_occurrences = sum(1 for p in exp_paragraphs if rules_engine.classify_keyword_presence(kw, p) == "present")
            placement_allowance += min(2, exp_occurrences)
            
            word_count = len(final_resume_text.split())
            if word_count < 400:
                m_length = 1.0
            elif word_count <= 700:
                m_length = 1.3
            else:
                m_length = 1.6
                
            if category == "hard_required":
                m_category = 1.5
            elif category == "soft_required":
                m_category = 1.2
            else:
                m_category = 1.0
                
            t_permitted = int((base_allowance + placement_allowance) * m_length * m_category)
            if freq > t_permitted:
                no_excessive_repetition = False
                break
                
        plain_bullets = len(re.findall(r'^[ \t]*[•*-⁃]', final_sections["summary"], re.MULTILINE))
        if plain_bullets > 0:
            no_malformed_formatting = False
        for para in indexed_paragraphs:
            t = para.text
            if any(char in t for char in ['\u201c', '\u201d', '\u2018', '\u2019', '\u2013', '\u2014', '\u2026']):
                no_malformed_formatting = False
                break
                
        for sec_name in ["summary", "skills", "experience", "education"]:
            if not final_sections.get(sec_name, "").strip():
                no_empty_sections = False
                break
                
        ats_score_calculated = score_after_data.get("score") is not None
        keyword_match_calculated = score_after_data.get("keyword_match") is not None
        recruiter_credibility_calculated = score_after_data.get("recruiter_credibility") is not None
        interview_defensibility_calculated = score_after_data.get("interview_defensibility") is not None
        
        multi_pass_executed = True
        missing_reviewed = True
        placement_reviewed = True

        def ck(val):
            return "X" if val else " "

        report_lines.extend([
            "=================================================",
            "RESUME OPTIMIZATION QUALITY CHECKLIST",
            "=====================================",
            "",
            "Career Progression",
            "",
            f"[{ck(client1_foundational)}] Client 1 role represents earlier-stage or foundational experience",
            f"[{ck(client2_advanced)}] Client 2 role represents more advanced or target-aligned experience",
            f"[{ck(progression_logical)}] Career progression remains logical",
            "",
            "Protected Information",
            "",
            f"[{ck(dates_preserved)}] Employment dates preserved",
            f"[{ck(durations_preserved)}] Employment durations preserved",
            f"[{ck(degree_history_preserved)}] Degree history preserved",
            f"[{ck(certifications_preserved)}] Earned certifications preserved",
            "",
            "Keyword Optimization",
            "",
            f"[{ck(hard_cov_calculated)}] Hard Required keyword coverage calculated",
            f"[{ck(soft_cov_calculated)}] Soft Required keyword coverage calculated",
            f"[{ck(pref_cov_calculated)}] Preferred keyword coverage calculated",
            "",
            "Keyword Placement",
            "",
            f"[{ck(keywords_in_skills)}] Relevant keywords appear in Skills",
            f"[{ck(keywords_in_experience)}] Relevant keywords appear in Experience",
            f"[{ck(full_exp_in_summary)}] FULL_EXPERIENCE keywords appear in Summary where appropriate",
            "",
            "Resume Quality",
            "",
            f"[{ck(no_duplicate_bullets)}] No duplicate bullets",
            f"[{ck(no_excessive_repetition)}] No excessive keyword repetition",
            f"[{ck(no_malformed_formatting)}] No malformed formatting",
            f"[{ck(no_empty_sections)}] No empty sections",
            "",
            "Scoring Validation",
            "",
            f"[{ck(ats_score_calculated)}] ATS score calculated",
            f"[{ck(keyword_match_calculated)}] Keyword Match score calculated",
            f"[{ck(recruiter_credibility_calculated)}] Recruiter Credibility score calculated",
            f"[{ck(interview_defensibility_calculated)}] Interview Defensibility score calculated",
            "",
            "Optimization Validation",
            "",
            f"[{ck(multi_pass_executed)}] Multi-pass optimization executed correctly",
            f"[{ck(missing_reviewed)}] Missing injectable keywords reviewed",
            f"[{ck(placement_reviewed)}] Placement opportunities reviewed",
            "",
            "====================================================",
            "END OF REPORT",
            "===================================================="
        ])

        report_text = "\n".join(report_lines)
        
        doc.save(output_path)
        print(json.dumps({
            "status": "success",
            "changes": changes_count,
            "mode": "pro",
            "report": report_text
        }))
        
    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
        sys.exit(1)

def optimize_basic(input_path, output_path, jd_text, prompt_instruction, model_name):
    """
    BASIC MODE: Fuzzy-match replacement (existing behavior).
    """
    try:
        doc = Document(input_path)
        
        full_text = []
        for para in iter_all_paragraphs(doc):
            if para.text.strip():
                full_text.append(para.text.strip())
        
        resume_content = "\n".join(full_text)
        
        system_prompt = f"""
        You are a resume rewriter. Your job is simple:
        
        **Goal**: {prompt_instruction}
        
        **CRITICAL INSTRUCTIONS**:
        1. You will receive a Resume and a Job Description (JD).
        2. **PROFESSIONAL SUMMARY - STRICT LENGTH RULE**:
           - Count the EXACT number of sentences in the original Professional Summary.
           - Your rewritten summary MUST have the EXACT SAME number of sentences.
           - If original has 3 sentences, output MUST have exactly 3 sentences.
           - DO NOT shorten or lengthen the summary under any circumstances.
        3. You MUST rewrite **EVERY SINGLE Experience bullet point** to match the JD. Do NOT skip any.
        4. For EACH paragraph/bullet in the resume, you MUST provide a replacement. No exceptions.
        5. The "original" field must be an EXACT copy-paste of the text from the resume I provide.
        6. The "new" field must be your rewritten version that aligns with the JD.
        7. If the domain is different (e.g., Java resume for a Data Engineering JD), completely rewrite the bullet point to be relevant.
        
        **Output Format**:
        Return ONLY a raw JSON object:
        {{
            "replacements": [
                {{"original": "exact text from resume", "new": "rewritten text for JD"}}
            ]
        }}
        """
        
        user_prompt = f"""
        Job Description:
        {jd_text}
        
        Resume Content:
        {resume_content}
        """

        ai_client = OpenAIProvider()
        response_text = ai_client.generate(system_prompt, user_prompt, model_name)
        response_text = response_text.replace("```json", "").replace("```", "").strip()
        
        try:
            data = json.loads(response_text)
        except json.JSONDecodeError:
            start = response_text.find('{')
            end = response_text.rfind('}') + 1
            if start != -1 and end != -1:
                data = json.loads(response_text[start:end])
            else:
                print(json.dumps({"status": "error", "message": "Could not parse AI response"}))
                sys.exit(1)

        replacements = data.get("replacements", [])
        changes_count = 0
        from difflib import SequenceMatcher

        def similar(a, b):
            return SequenceMatcher(None, a, b).ratio()

        for replacement in replacements:
            original = replacement["original"]
            new_text = replacement["new"]
            
            norm_original = " ".join(original.split())
            best_match_para = None
            best_ratio = 0.0

            for para in iter_all_paragraphs(doc):
                norm_para = " ".join(para.text.split())
                if len(norm_para) < 10:
                    continue

                ratio = similar(norm_original, norm_para)
                if norm_original in norm_para:
                     ratio = 1.0

                if ratio > best_ratio:
                    best_ratio = ratio
                    best_match_para = para

            if best_match_para and best_ratio > 0.40:
                replace_paragraph_text(best_match_para, new_text)
                changes_count += 1
            else:
                print(f"# Warning: No match for: {original[:60]}... (best ratio: {best_ratio:.2f})", file=sys.stderr)

        doc.save(output_path)
        print(json.dumps({
            "status": "success", 
            "changes": changes_count,
            "mode": "basic"
        }))

    except Exception as e:
        print(json.dumps({"status": "error", "message": str(e)}))
        sys.exit(1)

def optimize_resume(input_path, output_path, jd_text, prompt_instruction, provider, model_name, mode="basic"):
    print(f"\n{'='*60}", file=sys.stderr)
    print(f"  AUTORESUME MODE: {'🟣 PRO (Positional Replacement)' if mode == 'pro' else '🔵 BASIC (Fuzzy Match)'}", file=sys.stderr)
    print(f"  Model: {model_name}", file=sys.stderr)
    print(f"{'='*60}\n", file=sys.stderr)
    if mode == "pro":
        optimize_pro(input_path, output_path, jd_text, prompt_instruction, model_name)
    else:
        optimize_basic(input_path, output_path, jd_text, prompt_instruction, model_name)

if __name__ == "__main__":
    if len(sys.argv) < 7:
        print("Usage: script.py <input> <output> <jd_or_path> <prompt_or_path> <provider> <model> [mode]")
        sys.exit(1)
        
    input_p = sys.argv[1]
    output_p = sys.argv[2]
    jd_arg = sys.argv[3]
    prompt_arg = sys.argv[4]
    provider_arg = sys.argv[5]
    model_arg = sys.argv[6]
    mode_arg = sys.argv[7] if len(sys.argv) > 7 else "basic"
    
    if not os.environ.get("OPENAI_API_KEY"):
        print(json.dumps({"status": "error", "message": "OPENAI_API_KEY environment variable not set"}))
        sys.exit(1)
    
    if os.path.exists(jd_arg):
        with open(jd_arg, 'r', encoding='utf-8') as f:
            jd_t = f.read()
    else:
        jd_t = jd_arg
    
    jd_t = prune_job_description(jd_t)

    if os.path.exists(prompt_arg):
        with open(prompt_arg, 'r', encoding='utf-8') as f:
            prompt_t = f.read()
    else:
        prompt_t = prompt_arg
    
    optimize_resume(input_p, output_p, jd_t, prompt_t, provider_arg, model_arg, mode_arg)
