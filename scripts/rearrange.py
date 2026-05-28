"""
rearrange.py — Post-processing section rearrangement for AutoResume.

Takes an already-optimized DOCX (output from optimizer.py) and rearranges
its sections into a different order based on a template ID.

This script does NOT call any AI. It is pure document manipulation.
It preserves ALL formatting by moving actual XML elements, not copying text.

Usage:
    python rearrange.py <input.docx> <output.docx> <template_id>

Template IDs: standard, experience-led, balanced, executive
"""

import sys
import json
import os
import copy
from docx import Document
from lxml import etree

# ─── Section Detection ───────────────────────────────────────────────────────

SECTION_KEYWORDS = {
    "summary": ["professional summary", "summary", "career summary", "profile"],
    "skills":  ["technical skills", "skills", "core competencies", "technologies"],
    "experience": ["professional experience", "experience", "work experience", "employment"],
    "education": ["education", "academic", "qualifications"],
}

def classify_heading(text):
    """Check if a paragraph's text matches a known section heading."""
    lower = text.lower().strip()
    # Only classify short lines as headings (actual content paragraphs are longer)
    if len(lower) > 40:
        return None
    for section_key, keywords in SECTION_KEYWORDS.items():
        for kw in keywords:
            if kw in lower:
                return section_key
    return None

def is_heading_style(para):
    """Check if a paragraph uses a Word heading style (Heading 1, Heading 2, etc.)."""
    style_name = para.style.name if para.style else ""
    return style_name.startswith("Heading")

# ─── Core Rearrangement Logic ────────────────────────────────────────────────

def rearrange_sections(input_path, output_path, template_id):
    """
    Rearrange sections in the DOCX according to the template order.
    
    Strategy:
    1. Walk all body-level XML elements (paragraphs, tables, etc.)
    2. Detect section boundaries using Heading 2 style + keyword matching
    3. Group elements into section buckets
    4. Clear the document body
    5. Re-add elements in the template-specified order
    
    This moves actual lxml elements, so ALL formatting is preserved perfectly.
    """
    
    # Load template config
    templates_path = os.path.join(os.path.dirname(__file__), "templates.json")
    with open(templates_path, "r", encoding="utf-8") as f:
        templates = json.load(f)
    
    template = templates["templates"].get(template_id)
    if not template:
        print(json.dumps({"status": "error", "message": f"Unknown template: {template_id}"}))
        sys.exit(1)
    
    desired_order = template["order"]
    
    doc = Document(input_path)
    body = doc.element.body
    
    # Namespace for Word XML
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    
    # ─── Step 1: Collect all body-level elements and map paragraphs ───────
    # Body children include <w:p> (paragraphs), <w:tbl> (tables), 
    # <w:sectPr> (section properties), etc.
    all_elements = list(body)
    
    # Build a map: element -> paragraph object (only for <w:p> elements)
    # We need this to check styles and text
    para_map = {}
    for para in doc.paragraphs:
        para_map[para._element] = para
    
    # ─── Step 2: Walk elements and assign them to sections ────────────────
    # Sections are stored as: {"header": [elem, elem, ...], "summary": [...], ...}
    sections = {}
    current_section = "header"
    sections["header"] = []
    
    # Track sectPr (document section properties) — must stay at the end
    sect_pr_element = None
    
    for elem in all_elements:
        tag = etree.QName(elem.tag).localname if isinstance(elem.tag, str) else ""
        
        # sectPr must always be the very last element in the body
        if tag == "sectPr":
            sect_pr_element = elem
            continue
        
        # Check if this element is a paragraph we can inspect
        para = para_map.get(elem)
        
        if para is not None:
            text = para.text.strip()
            
            # Check for section heading: must match known section keywords
            # classify_heading already enforces len < 40 to avoid false positives
            if text:
                detected = classify_heading(text)
                if detected is not None:
                    current_section = detected
                    if current_section not in sections:
                        sections[current_section] = []
                    
                    # Standardize heading names and styles to match reference template
                    standard_names = {
                        "summary": "PROFESSIONAL SUMMARY",
                        "skills": "TECHNICAL SKILLS",
                        "experience": "WORK EXPERIENCE",
                        "education": "EDUCATION"
                    }
                    new_text = standard_names[detected]
                    if text != new_text:
                        para.text = ""
                        run = para.add_run(new_text)
                        try:
                            para.style = doc.styles['Heading 2']
                        except Exception:
                            try:
                                para.style = 'Heading 2'
                            except Exception:
                                pass
        
        sections.setdefault(current_section, []).append(elem)
    
    # Helper to clean leading and trailing empty paragraphs
    def clean_section_elements(elements):
        cleaned = list(elements)
        # Remove leading empty paragraphs
        while cleaned:
            first_elem = cleaned[0]
            para = para_map.get(first_elem)
            if para is not None and not para.text.strip():
                cleaned.pop(0)
            else:
                break
        # Remove trailing empty paragraphs
        while cleaned:
            last_elem = cleaned[-1]
            para = para_map.get(last_elem)
            if para is not None and not para.text.strip():
                cleaned.pop()
            else:
                break
        return cleaned

    # ─── Step 3: Clear the body and rebuild in desired order ──────────────
    # Remove all children from body
    for elem in all_elements:
        body.remove(elem)
    
    # Re-add elements in the template-specified order
    added_sections = set()
    for section_name in desired_order:
        if section_name in sections:
            cleaned_elements = clean_section_elements(sections[section_name])
            
            # Normalize skills content paragraphs if they were incorrectly styled as Heading 2
            if section_name == "skills":
                for idx, elem in enumerate(cleaned_elements):
                    if idx > 0:
                        para = para_map.get(elem)
                        if para is not None and para.style and para.style.name == "Heading 2":
                            try:
                                para.style = doc.styles['No Spacing']
                            except Exception:
                                try:
                                    para.style = 'No Spacing'
                                except Exception:
                                    try:
                                        para.style = 'Normal'
                                    except Exception:
                                        pass
            
            for elem in cleaned_elements:
                body.append(elem)
            added_sections.add(section_name)
    
    # Add any sections that weren't in the template order (safety net)
    for section_name, elements in sections.items():
        if section_name not in added_sections:
            print(f"# Warning: section '{section_name}' not in template order, appending at end", file=sys.stderr)
            cleaned_elements = clean_section_elements(elements)
            for elem in cleaned_elements:
                body.append(elem)
    
    # sectPr must always be the very last element
    if sect_pr_element is not None:
        body.append(sect_pr_element)
    
    # ─── Step 4: Clear bottom border on the last paragraph of the document (Balanced layout safety) ───
    if doc.paragraphs:
        last_para = doc.paragraphs[-1]
        if last_para.style and last_para.style.name == "Heading 2":
            try:
                last_para.style = 'No Spacing'
            except Exception:
                try:
                    last_para.style = 'Normal'
                except Exception:
                    pass
            
    # ─── Step 5: Save ─────────────────────────────────────────────────────
    doc.save(output_path)
    
    # Report results
    section_summary = {k: len(v) for k, v in sections.items()}
    print(json.dumps({
        "status": "success",
        "template": template_id,
        "sections_found": list(sections.keys()),
        "section_element_counts": section_summary,
        "reordered_to": desired_order
    }))


# ─── CLI Entry Point ─────────────────────────────────────────────────────────

if __name__ == "__main__":
    if len(sys.argv) < 4:
        print("Usage: python rearrange.py <input.docx> <output.docx> <template_id>")
        print("Template IDs: standard, experience-led, balanced, executive")
        sys.exit(1)
    
    input_p = sys.argv[1]
    output_p = sys.argv[2]
    template_arg = sys.argv[3]
    
    rearrange_sections(input_p, output_p, template_arg)
